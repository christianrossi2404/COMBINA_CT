
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement # Necesario para crear elementos XML como el salto de página
from docx.oxml.ns import qn
from pathlib import Path

# === FUNCIONES PARA CAMBIAR FUENTE ===

def cambiar_fuente_parrafos(doc):
    for parrafo in doc.paragraphs:
        for run in parrafo.runs:
            fuente = run.font
            # La condición para cambiar la fuente se mantiene como en tu código
            if fuente.name == "Times New Roman" and fuente.size == Pt(10):
                fuente.name = "Arial"
                fuente.size = Pt(9)
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.xpath('w:rFonts')
                for rf in rFonts:
                    rPr.remove(rf)
                rFonts_tag = rPr.get_or_add_rFonts()
                rFonts_tag.set(qn('w:ascii'), 'Arial')
                rFonts_tag.set(qn('w:hAnsi'), 'Arial')
                rFonts_tag.set(qn('w:cs'), 'Arial')

def cambiar_fuente_en_tablas(doc):
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        fuente = run.font
                        # La condición para cambiar la fuente se mantiene como en tu código
                        if fuente.name == "Times New Roman" and fuente.size == Pt(10):
                            fuente.name = "Arial"
                            fuente.size = Pt(9)
                            rPr = run._element.get_or_add_rPr()
                            rFonts = rPr.xpath('w:rFonts')
                            for rf in rFonts:
                                rPr.remove(rf)
                            rFonts_tag = rPr.get_or_add_rFonts()
                            rFonts_tag.set(qn('w:ascii'), 'Arial')
                            rFonts_tag.set(qn('w:hAnsi'), 'Arial')
                            rFonts_tag.set(qn('w:cs'), 'Arial')

# === INSERCIÓN DE SALTO DE PÁGINA (Función reintroducida) ===

def insertar_salto_pagina(doc):
    # Crea un nuevo elemento de párrafo XML
    salto_p = OxmlElement('w:p')
    # Crea un nuevo elemento de run XML dentro del párrafo
    run = OxmlElement('w:r')
    # Crea un elemento de salto de línea XML y establece su tipo a 'page'
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    # Añade el salto de línea a la run
    run.append(br)
    # Añade la run al párrafo
    salto_p.append(run)
    # Añade el párrafo (que contiene el salto de página) al cuerpo del documento
    doc.element.body.append(salto_p)
    print("   -> Salto de página insertado.")


# === FUNCION PARA COMBINAR DOCUMENTOS ===

def combinar_documentos_con_formato(plantilla_path, archivos_docx, salida_path):
    try:
        # Carga el documento base desde la plantilla proporcionada
        documento_base = Document(plantilla_path)
        print(f"✅ Plantilla '{plantilla_path.name}' cargada como documento base.")

        for i, archivo_path in enumerate(archivos_docx):
            try:
                # Abre cada documento a combinar
                doc = Document(archivo_path)
                print(f"🔄 Procesando '{archivo_path.name}'...")

                # Aplica los cambios de fuente al documento actual ANTES de combinarlo
                cambiar_fuente_parrafos(doc)
                cambiar_fuente_en_tablas(doc)
                print(f"   -> Fuente cambiada en '{archivo_path.name}'.")

                # Inserta un salto de página antes de cada documento, excepto el primero
                if i != 0:
                    insertar_salto_pagina(documento_base)

                # Añade todos los elementos del cuerpo del documento actual al documento base
                for elemento in doc.element.body:
                    documento_base.element.body.append(elemento)
                print(f"   -> Contenido de '{archivo_path.name}' añadido al documento base.")

            except FileNotFoundError:
                print(f"❌ Error: El archivo '{archivo_path.name}' no fue encontrado. Saltando este archivo.")
            except Exception as e:
                print(f"❌ Error al procesar '{archivo_path.name}': {e}")

        # Guarda el documento combinado final
        documento_base.save(salida_path)
        print(f"✅ Documentos combinados y guardados en: {salida_path}")

    except FileNotFoundError:
        print(f"❌ Error: La plantilla '{plantilla_path.name}' no fue encontrada. Asegúrate de que esté en el mismo directorio que el script.")
        # Es importante relanzar la excepción para que el programa no continúe con un error fatal.
        raise
    except Exception as e:
        print(f"❌ Error general durante la combinación de documentos: {e}")
        raise


# === USO DEL SCRIPT ===

if __name__ == "__main__":
    base_path = Path.cwd() # Obtiene el directorio de trabajo actual
    plantilla = base_path / "CT_TEMPLATE.docx" # Asegúrate de que esta plantilla exista
    archivos = [
        base_path / "01_CT-19782-20.docx", # Asegúrate de que estos archivos existan en el mismo directorio
        base_path / "02_CT-19782-21.docx"
    ]
    salida = base_path / "documento_combinado.docx"

    combinar_documentos_con_formato(plantilla, archivos, salida)