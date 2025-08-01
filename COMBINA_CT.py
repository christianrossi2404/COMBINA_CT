import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tkinter import Tk, filedialog, messagebox # ¡Añadir messagebox!

# === FUNCIONES PARA CAMBIAR FUENTE ===

def cambiar_fuente_run(run):
    """
    Cambia la fuente de un 'run' de texto de Times New Roman 10pt a Arial 9pt,
    incluyendo la manipulación XML para mayor robustez.
    """
    fuente = run.font
    if fuente.name == "Times New Roman" and fuente.size == Pt(10):
        fuente.name = "Arial"
        fuente.size = Pt(9)
        rPr = run._element.get_or_add_rPr()
        # Eliminar posibles elementos w:rFonts existentes para evitar conflicto
        for rf in rPr.xpath('w:rFonts'):
            rPr.remove(rf)
        # Agregar w:rFonts con la fuente deseada para distintos tipos de caracteres
        rFonts_tag = rPr.get_or_add_rFonts()
        rFonts_tag.set(qn('w:ascii'), 'Arial')
        rFonts_tag.set(qn('w:hAnsi'), 'Arial')
        rFonts_tag.set(qn('w:cs'), 'Arial')

def cambiar_fuente_parrafos(doc):
    """Aplica la función cambiar_fuente_run a todos los 'runs' en los párrafos del documento."""
    for parrafo in doc.paragraphs:
        for run in parrafo.runs:
            cambiar_fuente_run(run)

def cambiar_fuente_en_tablas(doc):
    """Aplica la función cambiar_fuente_run a todos los 'runs' dentro de las tablas del documento."""
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        cambiar_fuente_run(run)

# === INSERCIÓN DE SALTO DE PÁGINA ===

def insertar_salto_pagina(doc):
    """
    Inserta un salto de página explícito en el cuerpo del documento
    mediante la manipulación directa del XML.
    """
    salto_p = OxmlElement('w:p')
    run = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page') # Establece el tipo de salto a 'page'
    run.append(br)
    salto_p.append(run)
    doc.element.body.append(salto_p)
    # print("   -> Salto de página insertado.") # Este print no se verá en --windowed

# === COMBINACIÓN DE DOCUMENTOS ===

def combinar_documentos_con_formato(plantilla_path, archivos_docx, salida_path):
    """
    Combina múltiples documentos DOCX en uno solo, aplicando cambios de fuente
    y añadiendo saltos de página entre ellos (a partir del segundo).
    """
    try:
        # Carga el documento base desde la plantilla proporcionada
        documento_base = Document(plantilla_path)
        # print(f"✅ Plantilla '{plantilla_path.name}' cargada como documento base.") # No se verá

        for i, archivo_path in enumerate(archivos_docx):
            try:
                # Abre cada documento a combinar
                doc = Document(archivo_path)
                # print(f"🔄 Procesando '{archivo_path.name}'...") # No se verá

                # Aplica los cambios de fuente al documento actual ANTES de combinarlo
                cambiar_fuente_parrafos(doc)
                cambiar_fuente_en_tablas(doc)
                # print(f"   -> Fuente cambiada en '{archivo_path.name}'.") # No se verá

                # Inserta un salto de página antes de cada documento, excepto el primero
                if i != 0:
                    insertar_salto_pagina(documento_base)

                # Añade todos los elementos del cuerpo del documento actual al documento base
                for elemento in doc.element.body:
                    documento_base.element.body.append(elemento)
                # print(f"   -> Contenido de '{archivo_path.name}' añadido al documento base.") # No se verá

            except FileNotFoundError:
                messagebox.showerror("Error de Archivo", f"El archivo '{archivo_path.name}' no fue encontrado. Saltando este archivo.")
            except Exception as e:
                messagebox.showerror("Error de Procesamiento", f"Error al procesar '{archivo_path.name}': {e}")

        # Guarda el documento combinado final
        documento_base.save(salida_path)
        messagebox.showinfo("Proceso Completado", f"Documentos combinados y guardados en:\n{salida_path}")

    except FileNotFoundError:
        messagebox.showerror("Error de Plantilla", f"La plantilla '{plantilla_path.name}' no fue encontrada. Asegúrate de que esté en el mismo directorio que el script.")
        raise
    except Exception as e:
        messagebox.showerror("Error General", f"Error general durante la combinación de documentos: {e}")
        raise

# === DETECCIÓN Y SELECCIÓN DE ARCHIVOS DOCX ===

def obtener_archivos_desde_argumentos():
    return [Path(f) for f in sys.argv[1:] if f.lower().endswith('.docx') and Path(f).exists()]

def seleccionar_archivos_manual():
    root = Tk()
    root.withdraw() # Ocultar la ventana principal de Tkinter
    archivos = filedialog.askopenfilenames(
        title="Selecciona los archivos DOCX a combinar (puedes seleccionar varios)",
        filetypes=[("Documentos Word", "*.docx")]
    )
    root.destroy()
    return [Path(f) for f in archivos]

# === EJECUCIÓN PRINCIPAL DEL SCRIPT ===

if __name__ == "__main__":
    # Inicializar Tkinter para que los cuadros de diálogo funcionen correctamente
    # incluso si no se abre el selector de archivos (por ejemplo, al usar "Send To")
    root_tk_main = Tk()
    root_tk_main.withdraw() # Ocultar la ventana principal de Tkinter

    archivos_docx = obtener_archivos_desde_argumentos()

    if not archivos_docx:
        # print("No se detectaron archivos pasados como argumento. Abriendo selector manual...") # No se verá
        archivos_docx = seleccionar_archivos_manual()

    if not archivos_docx:
        messagebox.showwarning("Sin Selección", "No se seleccionaron archivos. El programa finalizará.")
        root_tk_main.destroy() # Destruir la instancia de Tkinter
        sys.exit()

    # Determinar la carpeta de destino: la misma que la del primer archivo seleccionado
    carpeta_destino = archivos_docx[0].parent
    salida = carpeta_destino / "DOCUMENTO_COMBINADO.docx" # Nombre del archivo de salida

    # Ruta a la plantilla base
    plantilla = Path(__file__).parent / "CT_TEMPLATE.docx"

    # Llamar a la función principal para combinar los documentos
    combinar_documentos_con_formato(plantilla, archivos_docx, salida)

    root_tk_main.destroy() # Destruir la instancia de Tkinter al finalizar el proceso
    # input("\nProceso completado. Presiona Enter para salir...") # ¡ELIMINADO!